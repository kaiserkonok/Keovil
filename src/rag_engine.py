import os
import json
import shutil
from datetime import datetime
import threading
import pdfplumber
import docx
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import OllamaEmbeddings
from langchain_ollama import OllamaLLM
from sentence_transformers import CrossEncoder
import intelligent_rag_chunker
from pathlib import Path
import torch
import numpy as np
from rank_bm25 import BM25Okapi
from typing import List, Dict, Any, Union
from concurrent.futures import ThreadPoolExecutor


# ----------------------
# ANSI Color Class for Debugging
# ----------------------
class Colors:
    HEADER = '\033[95m'  # Magenta
    OKBLUE = '\033[94m'  # Blue
    OKCYAN = '\033[96m'  # Cyan
    OKGREEN = '\033[92m'  # Green
    WARNING = '\033[93m'  # Yellow
    FAIL = '\033[91m'  # Red
    ENDC = '\033[0m'  # Reset color
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'


# ----------------------
# RRF Utility Function (Outside the Class)
# ----------------------
RRF_K = 60


def reciprocal_rank_fusion(
        results: List[List[Document]],
        k: int = RRF_K
) -> List[Document]:
    """
    Combines ranked results from multiple lists (e.g., Vector and BM25)
    using Reciprocal Rank Fusion (RRF).
    """
    fused_scores = {}
    doc_map = {}

    # Iterate over each retrieval method's results
    for rank_list in results:
        # doc_id is typically a unique hash or combination of source/chunk_id
        for rank, doc in enumerate(rank_list):
            # Ensure doc is a Document object before accessing metadata
            if not isinstance(doc, Document):
                continue

            doc_id = doc.metadata.get("source") + str(doc.metadata.get("chunk_id"))

            # Map doc_id to Document object
            if doc_id not in doc_map:
                doc_map[doc_id] = doc

            # Calculate RRF score: 1 / (K + rank + 1)
            score = 1 / (k + rank + 1)

            # Accumulate the score
            fused_scores[doc_id] = fused_scores.get(doc_id, 0) + score

    # Sort document IDs by the fused RRF score
    sorted_doc_ids = sorted(fused_scores, key=fused_scores.get, reverse=True)

    # Return the final list of Document objects
    return [doc_map[doc_id] for doc_id in sorted_doc_ids]


# ----------------------
# Watchdog handler for real-time ingestion
# ----------------------
class NewFileHandler(FileSystemEventHandler):
    def __init__(self, rag_instance):
        self.rag = rag_instance

    def on_created(self, event):
        if event.is_directory:
            return
        print(f"{Colors.OKCYAN}[Watcher]{Colors.ENDC} New file detected: {event.src_path}")
        self.rag.ingest(new_files=[event.src_path])

    def on_modified(self, event):
        if event.is_directory:
            return
        print(f"{Colors.OKCYAN}[Watcher]{Colors.ENDC} File modified: {event.src_path}")
        self.rag.ingest(new_files=[event.src_path])

    def on_deleted(self, event):
        if event.is_directory:
            return
        print(f"{Colors.WARNING}[Watcher]{Colors.ENDC} File deleted: {event.src_path}")
        self.rag.remove_file(event.src_path)


# ----------------------
# College RAG System
# ----------------------
class CollegeRAG:
    def __init__(self, data_dir=None, top_k=7, rerank_top_k=15, store_dir=None, llm_model='qwen2.5:14b-instruct-q4_K_M', temperature=0):
        # --- PATH CONFIGURATION (Home Directory Defaults) ---
        home = str(Path.home())
        base_storage = os.path.join(home, ".k_rag_storage")

        if data_dir is None:
            data_dir = os.path.join(base_storage, "data")
        if store_dir is None:
            store_dir = os.path.join(base_storage, "faiss_store")

        self.data_dir = os.path.abspath(data_dir)
        self.store_dir = os.path.abspath(store_dir)

        # Ensure directories exist
        os.makedirs(self.data_dir, exist_ok=True)
        os.makedirs(self.store_dir, exist_ok=True)

        self.top_k = top_k
        self.rerank_top_k = rerank_top_k
        self.docs = []
        # self.chat_history is used only for the interactive terminal test
        self.chat_history = []
        self.vectorestore = None
        self.bm25_retriever = None
        self.doc_map = {}
        self.lock = threading.Lock()

        # LLM + reranker (MODIFIED: Now accepts model and temperature dynamically)
        self.llm = OllamaLLM(model=llm_model, streaming=True, temperature=temperature)

        # Intelligence LLM (3B) for query strategy
        self.intel_llm = OllamaLLM(model='qwen2.5:14b-instruct-q4_K_M', temperature=0)

        self.embedding_model = OllamaEmbeddings(model='bge-m3:latest')

        device = 'cuda' if torch.cuda.is_available() else 'cpu'
        print(f"{Colors.HEADER}Reranker is initializing on device: {Colors.BOLD}{device}{Colors.ENDC}")

        # --- SAFE INITIALIZATION ---
        # Upgraded to BGE Reranker for better accuracy
        self.reranker = CrossEncoder(
            model_name='cross-encoder/ms-marco-MiniLM-L-6-v2',
            device=device
        )

        # Initialize your intelligent chunker
        self.chunker = intelligent_rag_chunker.IntelligentChunker()

        # Load docs and create vectorstore
        self.create_vectorestore(force_create=False)

        if self.vectorestore:
            # Rebuild self.docs and BM25 index from loaded store
            self.docs = list(self.vectorestore.docstore._dict.values())
            self._build_bm25_index()

        # Start Watchdog for automatic ingestion
        event_handler = NewFileHandler(self)
        self.observer = Observer()
        self.observer.schedule(event_handler, self.data_dir, recursive=True)
        self.observer.start()

    # ----------------------
    # Helper to build/rebuild the BM25 index
    # ----------------------
    def _build_bm25_index(self):
        if not self.docs:
            self.bm25_retriever = None
            self.doc_map = {}
            return

        print(f'{Colors.OKBLUE}[Hybrid Search] Creating BM25 index...{Colors.ENDC}')

        # 1. Prepare tokenized corpus
        bm25_tokenized_corpus = [
            doc.page_content.lower().split(" ") for doc in self.docs
        ]

        # 2. Initialize the BM25 retriever
        self.bm25_retriever = BM25Okapi(bm25_tokenized_corpus)

        # 3. Create map from corpus index to Document object
        self.doc_map = {i: doc for i, doc in enumerate(self.docs)}

        print(f"{Colors.OKBLUE}[Hybrid Search] BM25 index created with {len(self.docs)} documents.{Colors.ENDC}")

    # ----------------------
    # Load and chunk documents using IntelligentChunker
    # ----------------------
    def _load_docs(self):
        self.docs = self.get_docs(self.data_dir)

    def get_docs(self, data_dir, docs=None):
        if docs is None:
            docs = []

        # FIX: Check if directory exists before listing
        if not os.path.exists(data_dir):
            os.makedirs(data_dir, exist_ok=True)
            return docs

        for fname in os.listdir(data_dir):
            fpath = os.path.join(data_dir, fname)
            if os.path.isdir(fpath):
                self.get_docs(fpath, docs)
                continue

            ext = fname.lower().split('.')[-1]
            try:
                text = ""
                if ext == 'txt':
                    with open(fpath, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()

                elif ext == 'docx':
                    doc = docx.Document(fpath)
                    text = "\n".join([p.text for p in doc.paragraphs])

                elif ext == 'pdf':
                    with pdfplumber.open(fpath) as pdf:
                        pages = []
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                pages.append(page_text)
                        text = "\n".join(pages)

                if text:
                    # Use your intelligent chunker
                    chunk_objs = self.chunker.chunk_document(text)
                    for c in chunk_objs:
                        docs.append(Document(
                            page_content=c.text,
                            metadata={"source": fpath, "chunk_id": c.id, **c.meta}
                        ))

            except Exception as e:
                print(f"{Colors.FAIL}Failed {fpath}: {e}{Colors.ENDC}")
                continue

        return docs

    # ----------------------
    # Create or load FAISS vectorstore
    # ----------------------
    def create_vectorestore(self, force_create=False):
        with self.lock:
            store_path = self.store_dir

            if force_create:
                print(f"{Colors.WARNING}Force Create active. Wiping old index and clearing memory...{Colors.ENDC}")
                self.vectorestore = None
                self.docs = []  # <--- Crucial: Wipes the in-memory list
                self.bm25_retriever = None

                if os.path.exists(store_path):
                    # We only delete store_dir (the FAISS cache), NOT data_dir (your PDFs)
                    shutil.rmtree(store_path)
                os.makedirs(store_path, exist_ok=True)

            # Check if the folder is empty or non-existent
            store_exists = os.path.exists(store_path) and len(os.listdir(store_path)) > 0

            if store_exists and not force_create:
                try:
                    print(f'{Colors.OKGREEN}Loading existing FAISS vectorstore from {store_path}...{Colors.ENDC}')
                    self.vectorestore = FAISS.load_local(
                        store_path,
                        self.embedding_model,
                        allow_dangerous_deserialization=True
                    )
                except Exception as e:
                    print(f"{Colors.WARNING}Failed to load store: {e}. Recreating...{Colors.ENDC}")
                    self.vectorestore = None

            # Rebuild if forced or if loading failed
            if self.vectorestore is None:
                print(f'{Colors.WARNING}Scanning source files for fresh indexing...{Colors.ENDC}')
                # Since self.docs was cleared above, get_docs will only find current files
                self._load_docs()

                if self.docs:
                    self.vectorestore = FAISS.from_documents(self.docs, self.embedding_model)
                    self.vectorestore.save_local(store_path)
                    self._build_bm25_index()
                else:
                    print(f"{Colors.FAIL}No documents found in {self.data_dir}.{Colors.ENDC}")

            # Ensure the retriever is updated to use the new index
            if self.vectorestore:
                self.retriever = self.vectorestore.as_retriever(search_kwargs={'k': self.rerank_top_k})

    # ----------------------
    # Incremental ingestion
    # ----------------------
    def ingest(self, new_files=None):
        """
        Incrementally ingest files and rebuilds the BM25 index afterwards.
        """
        with self.lock:
            if not new_files:
                return

            all_new_docs = []  # Collect all new chunks

            for fpath in new_files:
                file_path = Path(fpath)

                if file_path.name.endswith("~") or file_path.name.startswith(".#") or file_path.suffix == ".swp":
                    print(f"{Colors.WARNING}[Ingest] Skipping temporary file: {fpath}{Colors.ENDC}")
                    continue

                print(f"{Colors.BOLD}[Ingest] Updating file: {fpath}{Colors.ENDC}")

                # Remove old chunks only if vectorestore exists
                if self.vectorestore:
                    ids_to_delete = [k for k, doc in self.vectorestore.docstore._dict.items()
                                     if doc.metadata.get("source") == fpath]

                    if ids_to_delete:
                        self.vectorestore.delete(ids_to_delete)
                        print(f"{Colors.OKBLUE}[Ingest] Removed {len(ids_to_delete)} old chunks.{Colors.ENDC}")

                ext = file_path.suffix.lower()[1:]
                try:
                    if ext == "txt":
                        with open(fpath, 'r', encoding='utf-8', errors='ignore') as f:
                            text = f.read()
                    elif ext == "docx":
                        docx_doc = docx.Document(fpath)
                        text = "\n".join([p.text for p in docx_doc.paragraphs])
                    elif ext == "pdf":
                        with pdfplumber.open(fpath) as pdf:
                            pages = [p.extract_text() for p in pdf.pages if p.extract_text()]
                            text = "\n".join(pages)
                    else:
                        print(f"{Colors.FAIL}[Ingest] Unsupported file type: {fpath}{Colors.ENDC}")
                        continue
                except Exception as e:
                    print(f"{Colors.FAIL}[Ingest] Failed to read {fpath}: {e}{Colors.ENDC}")
                    continue

                if not text.strip():
                    print(f"{Colors.WARNING}[Ingest] File {fpath} is empty. Skipping.{Colors.ENDC}")
                    continue

                chunks = self.chunker.chunk_document(text.strip())
                new_file_docs = [
                    Document(page_content=c.text, metadata={"source": fpath, "chunk_id": c.id, **c.meta})
                    for c in chunks
                ]
                all_new_docs.extend(new_file_docs)

            if not all_new_docs:
                return

            # FIX: If vectorestore was None (empty start), create it now. Else, add.
            if self.vectorestore is None:
                self.vectorestore = FAISS.from_documents(all_new_docs, self.embedding_model)
            else:
                self.vectorestore.add_documents(all_new_docs)

            print(f"{Colors.OKGREEN}[Ingest] Added {len(all_new_docs)} new chunks.{Colors.ENDC}")

            self.docs = list(self.vectorestore.docstore._dict.values())
            self._build_bm25_index()
            self.retriever = self.vectorestore.as_retriever(search_kwargs={'k': self.rerank_top_k})

            self.vectorestore.save_local(self.store_dir)
            print(f"{Colors.OKGREEN}[Ingest] Ingestion completed!{Colors.ENDC}")

    def remove_file(self, fpath):
        """
        Remove all chunks of a deleted file from the vectorstore and BM25 index.
        """
        with self.lock:
            if not self.vectorestore:
                return

            ids_to_delete = [
                doc_id for doc_id, doc in self.vectorestore.docstore._dict.items()
                if doc.metadata.get("source") == fpath
            ]

            if not ids_to_delete:
                print(f"{Colors.WARNING}[Remove] No chunks found for deleted file: {fpath}{Colors.ENDC}")
                return

            self.vectorestore.delete(ids_to_delete)
            print(f"{Colors.OKBLUE}[Remove] Deleted {len(ids_to_delete)} chunks from: {fpath}{Colors.ENDC}")

            # Update docs and rebuild BM25
            self.docs = list(self.vectorestore.docstore._dict.values())

            # FIX: If last file deleted, reset store to None
            if not self.docs:
                self.vectorestore = None
                self.bm25_retriever = None
                print(f"{Colors.WARNING}[Remove] All documents removed. Store is empty.{Colors.ENDC}")
            else:
                self._build_bm25_index()
                self.retriever = self.vectorestore.as_retriever(search_kwargs={'k': self.rerank_top_k})
                self.vectorestore.save_local(self.store_dir)

            print(f"{Colors.OKGREEN}[Remove] Vectorstore updated.{Colors.ENDC}")

    # ----------------------
    # Get Search Intelligence (World Class Upgrade)
    # ----------------------
    def get_search_intelligence(self, query, chat_history):
        prompt = f"""
            You are a Search Intent Architect. Your goal is to prepare a multi-path search strategy.

            ### TASK 1: STANDALONE QUERY
            Rewrite the "User Query" to be entirely self-contained. 
            
            - Keep the query minimal. Do not add extra words, names, or details unless they are strictly required to resolve pronouns or ambiguity.
            - If the query is already clear, leave it unchanged.
            - Resolve pronouns, fragments, or vague references using the chat history so the query stands alone.
            - Do not expand the query with institution names, locations, or other context unless the user explicitly mentioned them.
            - Your rewritten query should be as short and focused as possible.

            ### TASK 2: HYPOTHETICAL DOCUMENT (HyDE)
            Create a brief snippet of text that would serve as a perfect, direct answer to the standalone query. 
            - Use professional and factual language.
            - This is for semantic matching; focus on how a document would logically state the information.

            ### TASK 3: KEYWORD EXTRACTION
            Extract the most significant search terms strictly from the "Standalone Query".
            - Focus on unique names, technical nouns, and core actions.
            - Do not include common stop words or words only found in the HyDE section.

            ### CONTEXT:
            Chat History: {chat_history}
            User Query: {query}

            ### OUTPUT FORMAT:
            JSON only.
            {{
                "standalone": "string",
                "hyde": "string",
                "keywords": ["list", "of", "terms"]
            }}
        """.strip()

        try:
            raw_response = self.llm.invoke(prompt)
            start = raw_response.find('{')
            end = raw_response.rfind('}') + 1
            if start == -1: return {"standalone": query, "hyde": query, "keywords": query.split()}

            intel = json.loads(raw_response[start:end])

            # Smart Filtering: Ensure keywords exist in the standalone text
            intel['keywords'] = [w for w in intel['keywords'] if w.lower() in intel['standalone'].lower()]

            return intel
        except Exception:
            return {"standalone": query, "hyde": query, "keywords": [w for w in query.split() if len(w) > 3]}

    # ----------------------
    # Ask a question (streaming or normal) - MODIFIED FOR HISTORY MANAGEMENT
    # ----------------------
    def ask(self, query, chat_history: Union[List[tuple], None] = None, stream=False):
        use_internal_history = False
        if chat_history is None:
            chat_history = self.chat_history
            use_internal_history = True

        print(f"\n{Colors.HEADER}{'=' * 20} RAG PIPELINE START {'=' * 20}{Colors.ENDC}")

        # 1️⃣ Fast Intent Extraction (Intelligence)
        print(f"{Colors.OKCYAN}[1/5] Analyzing Intent...{Colors.ENDC}")
        intel = self.get_search_intelligence(query, chat_history)
        rewritten_query = intel.get('standalone', query)
        keywords = intel.get('keywords', [])

        print(f"    {Colors.BOLD}• Rewritten:{Colors.ENDC} {rewritten_query}")
        print(f"    {Colors.BOLD}• Keywords:{Colors.ENDC} {', '.join(keywords)}")
        print(f"    {Colors.BOLD}• Hyde:{Colors.ENDC} {intel['hyde']}")

        if not self.vectorestore:
            print(f"{Colors.FAIL}[Error] Vectorstore is empty!{Colors.ENDC}")
            return "I haven't learned anything yet!"

        # 2️⃣ Parallel Multi-Path Retrieval
        print(f"{Colors.OKCYAN}[2/5] Executing Hybrid Retrieval (Vector + BM25)...{Colors.ENDC}")
        k_initial = 25

        def vector_search(q):
            return self.vectorestore.similarity_search(q, k=k_initial)

        def bm25_search(kw):
            tokenized_query = " ".join(kw).lower().split()
            scores = self.bm25_retriever.get_scores(tokenized_query)
            indices = np.argsort(scores)[::-1][:k_initial]
            return [self.doc_map[i] for i in indices if i in self.doc_map]

        with ThreadPoolExecutor(max_workers=3) as executor:
            f1 = executor.submit(vector_search, rewritten_query)
            f2 = executor.submit(vector_search, intel.get('hyde', rewritten_query))
            f3 = executor.submit(bm25_search, keywords)
            results = [f1.result(), f2.result(), f3.result()]

        # 3️⃣ Deduplication & RRF
        fused_candidates = reciprocal_rank_fusion(results)
        print(
            f"    {Colors.OKBLUE}• Retrieved {sum(len(r) for r in results)} raw chunks. Fused into {len(fused_candidates)} unique candidates.{Colors.ENDC}")

        # 4️⃣ Pruned Reranking (GPU Accelerated)
        print(f"{Colors.OKCYAN}[3/5] Reranking top 20 candidates on RTX 5060 Ti...{Colors.ENDC}")
        rerank_pool = fused_candidates[:20]
        pairs = [[rewritten_query, doc.page_content] for doc in rerank_pool]

        # Calculate scores
        scores = self.reranker.predict(pairs, batch_size=20, show_progress_bar=False)
        ranked_indices = np.argsort(scores)[::-1]

        # Debug: Show the top 3 match scores
        top_scores = [round(float(scores[i]), 4) for i in ranked_indices[:3]]
        print(f"    {Colors.OKGREEN}• Top Relevancy Scores: {top_scores}{Colors.ENDC}")

        # 5️⃣ Dynamic Selection
        final_top_k = 6
        top_docs = [rerank_pool[i] for i in ranked_indices[:final_top_k]]

        print(f"{Colors.OKCYAN}[4/5] Context prepared from {len(top_docs)} sources.{Colors.ENDC}")
        for i, doc in enumerate(top_docs[:2]):  # Show snippet of top 2 for debugging
            src = os.path.basename(doc.metadata.get('source', 'unknown'))
            snippet = doc.page_content[:75].replace('\n', ' ')
            print(f"      {Colors.OKBLUE}[Source {i + 1}]{Colors.ENDC} {src}: \"{snippet}...\"")

        # 6️⃣ Stream/Generate Response
        print(f"{Colors.OKCYAN}[5/5] Generating Answer...{Colors.ENDC}")
        print(f"{Colors.HEADER}{'=' * 56}{Colors.ENDC}\n")

        context_text = "\n\n".join([d.page_content for d in top_docs])
        prompt = f"""Lora, use context to answer.
        HISTORY: {chat_history}
        CONTEXT: {context_text}
        QUESTION: {query}
        ANSWER:"""

        if stream:
            response = ""
            print(f"{Colors.BOLD}Lora:{Colors.ENDC} ", end="", flush=True)
            for chunk in self.llm.stream(prompt):
                print(chunk, end="", flush=True)
                response += chunk
            if use_internal_history:
                self.chat_history.append(("You", query))
                self.chat_history.append(("AI", response))
            print(f"\n\n{Colors.HEADER}{'=' * 56}{Colors.ENDC}")
            return response

        res = self.llm.invoke(prompt)
        if use_internal_history:
            self.chat_history.append(("You", query))
            self.chat_history.append(("AI", res))
        return res


# ----------------------
# Run interactively
# ----------------------
if __name__ == "__main__":
    # Updated to use home-based default storage if no data_dir provided
    rag = CollegeRAG()
    rag.create_vectorestore(force_create=True)

    if rag.docs:
        for doc in rag.docs:
            print(
                f'{Colors.HEADER}___ Starting of Chunk ({os.path.basename(doc.metadata.get("source", "N/A"))}) ___{Colors.ENDC}\n')
            print(doc.page_content)
            print(f'\n{Colors.HEADER}___ Ending of Chunk ___{Colors.ENDC}')
    else:
        print(f"{Colors.WARNING}No initial documents found. Waiting for uploads to {rag.data_dir}...{Colors.ENDC}")

    try:
        while True:
            question = input(f"{Colors.BOLD}You:{Colors.ENDC} ")
            # History works here because rag.ask uses and updates self.chat_history
            rag.ask(question, stream=True)
    except KeyboardInterrupt:
        print(f"\n{Colors.FAIL}Stopping...{Colors.ENDC}")
        rag.observer.stop()
        rag.observer.join()